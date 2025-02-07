from django.shortcuts import render,redirect
from django.http import HttpResponse, FileResponse
import img2pdf
import fitz
from io import BytesIO
import os
from django.conf import settings
from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# Create your views here.
def Main(request):
    return render(request, 'main.html')



def ImageToPdf(request):
    if request.method == "POST" and request.FILES.getlist('image'):
        image_files = request.FILES.getlist('image')

        pdf_data = img2pdf.convert([img.read() for img in image_files])

        # Without download button
        # response = HttpResponse(pdf_data, content_type="application/pdf")
        # response['Content-Disposition'] = 'attachment; filename="FlexiFile.pdf"'
        # return response

        pdf_filename = 'FlexiFileI2P.pdf'
        pdf_path = os.path.join(settings.MEDIA_ROOT, pdf_filename)

        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)  



        with open(pdf_path, "wb")as f:
            f.write(pdf_data)

        return render(request, 'imgtopdf.html', {'pdf_filename':pdf_filename})

    return render(request, 'imgtopdf.html')


def download_pdf(request, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    
    if os.path.exists(file_path):
        return FileResponse(open(file_path, "rb"), as_attachment=True, content_type="application/pdf")

    return HttpResponse("File not found!!", status=404)






def PdfToImage(request):
    if request.method == "POST" and request.FILES.get('pdf'):
        pdf_file = request.FILES['pdf']

        # Ensure the media directory exists
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)

        # Save uploaded PDF
        pdf_filename = "uploaded.pdf"
        pdf_path = os.path.join(settings.MEDIA_ROOT, pdf_filename)
        with open(pdf_path, "wb") as f:
            f.write(pdf_file.read())

        # Open PDF and convert each page to image
        doc = fitz.open(pdf_path)
        image_filenames = []

        for i, page in enumerate(doc):
            pix = page.get_pixmap()
            image_filename = f"FlexiFilepage_{i + 1}.png"
            image_path = os.path.join(settings.MEDIA_ROOT, image_filename)
            print(f"Image saved at: {image_path}")
            pix.save(image_path)  # Save the page as PNG

            # Save image filenames for rendering in the template
            image_filenames.append(image_filename)

        return render(request, 'pdftoimg.html', {'image_filenames': image_filenames})

    return render(request, 'pdftoimg.html')






def PdfToWord(request):
    if request.method == "POST" and request.FILES.get('pdf'):
        pdf_file = request.FILES['pdf']

        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)

        pdf_filename = "uploaded.pdf"
        pdf_path = os.path.join(settings.MEDIA_ROOT, pdf_filename)
        with open(pdf_path, "wb") as f:
            f.write(pdf_file.read())

        word_filename = 'FlexiFileP2W.docx'
        word_path = os.path.join(settings.MEDIA_ROOT, word_filename)

        try:
            pdf_reader = PdfReader(pdf_path)
            doc = Document()

            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
                else:
                    doc.add_paragraph("[Image or Unreadable Page]")  

            doc.save(word_path)
        except Exception as e:
            return HttpResponse(f"Error: {str(e)}", status=500)
        
        return render(request, 'pdftoword.html', {'word_filename': word_filename})

    return render(request, 'pdftoword.html')


def download_word(request, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    
    if os.path.exists(file_path):
        return FileResponse(open(file_path, "rb"), as_attachment=True, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return HttpResponse("File not found!!", status=404)






def WordToPdf(request):
    if request.method == "POST" and request.FILES.get('word'):
        word_file = request.FILES['word']

        # Ensure the media directory exists
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)

        # Save the uploaded Word file
        word_filename = "uploaded.docx"
        word_path = os.path.join(settings.MEDIA_ROOT, word_filename)
        with open(word_path, "wb") as f:
            f.write(word_file.read())

        # Convert Word to PDF
        pdf_filename = "FlexiFileW2P.pdf"
        pdf_path = os.path.join(settings.MEDIA_ROOT, pdf_filename)

        try:
            doc = Document(word_path)
            pdf_canvas = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4

            y_position = height - 50  # Start position for text

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    pdf_canvas.drawString(50, y_position, text)
                    y_position -= 20  # Move down for next line

                    # Page break handling
                    if y_position < 50:
                        pdf_canvas.showPage()
                        y_position = height - 50

            pdf_canvas.save()
        except Exception as e:
            return HttpResponse(f"Error: {str(e)}", status=500)

        return render(request, 'wordtopdf.html', {'pdf_filename': pdf_filename})

    return render(request, 'wordtopdf.html')






def PdfMerge(request):
    if request.method == "POST" and request.FILES.getlist('pdfs'):
        pdf_files = request.FILES.getlist('pdfs')

        # Ensure the media directory exists
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)

        merged_filename = "FlexiFilemerged.pdf"
        merged_path = os.path.join(settings.MEDIA_ROOT, merged_filename)

        pdf_writer = PdfWriter()

        try:
            for pdf in pdf_files:
                pdf_path = os.path.join(settings.MEDIA_ROOT, pdf.name)
                
                # Save each PDF temporarily
                with open(pdf_path, "wb") as f:
                    f.write(pdf.read())

                # Append to PdfWriter
                pdf_writer.append(pdf_path)

            # Write merged PDF
            with open(merged_path, "wb") as output_pdf:
                pdf_writer.write(output_pdf)

        except Exception as e:
            return HttpResponse(f"Error: {str(e)}", status=500)

        return render(request, 'pdfmerge.html', {'merged_filename': merged_filename})

    return render(request, 'pdfmerge.html')






def PdfCompress(request):
    if request.method == "POST" and request.FILES.get('pdf'):
        pdf_file = request.FILES['pdf']

        # Ensure the media directory exists
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)

        # Save uploaded PDF
        pdf_filename = "uploaded.pdf"
        pdf_path = os.path.join(settings.MEDIA_ROOT, pdf_filename)
        with open(pdf_path, "wb") as f:
            f.write(pdf_file.read())

        # Define compressed PDF path
        compressed_filename = "FlexiFilecompressed.pdf"
        compressed_path = os.path.join(settings.MEDIA_ROOT, compressed_filename)

        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()

            for page in reader.pages:
                writer.add_page(page)  # ✅ Add the page first
                writer.pages[-1].compress_content_streams()  # ✅ Then compress it

            # Save compressed PDF
            with open(compressed_path, "wb") as output_pdf:
                writer.write(output_pdf)

        except Exception as e:
            return HttpResponse(f"Error: {str(e)}", status=500)

        return render(request, 'pdfcompress.html', {'compressed_filename': compressed_filename})

    return render(request, 'pdfcompress.html')






from .models import ContactSubmission
from django.contrib import messages

#contact (get in touch)
def contact_view(request):
    if request.method == "POST":
        name = request.POST.get('name')
        email = request.POST.get('email')
        message = request.POST.get('message')

        ContactSubmission.objects.create(name=name, email = email, message=message)

        messages.success(request, "Your message has been submitted successfully!!")
        return redirect('contact')
    
    return render(request, 'main.html')
